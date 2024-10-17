import streamlit as st
import json
from docx import Document
import io
import os 
import streamlit as st 
import re 
import json
from langchain.prompts import PromptTemplate
from docx import Document
from dotenv import load_dotenv
import tempfile
from langchain_google_genai import ChatGoogleGenerativeAI
import json
import io
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.shared import Pt

load_dotenv()
Gemini = st.secrets["GOOGLE_API_KEY"]
llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro-latest", google_api_key=Gemini)

# The UI interface 
st.set_page_config(page_title="LEGAL DRAFT PRO", page_icon="⚖️", layout="wide")
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;700&display=swap');
    
    .stApp {
        background-color: #f0f0f0;
        font-family: 'Roboto', sans-serif;
    }
    .container {
        max-width: 1000px;
        margin: 0 auto;
        padding: 20px;
        background-color: white;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        border-radius: 10px;
    }
    .big-font {
        font-size: 48px !important;
        font-weight: 700;
        color: #6a0dad;
        text-align: center;
        margin-bottom: 10px;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
    }
    .medium-font {
        font-size: 24px !important;
        font-weight: 400;
        color: #008080;
        text-align: center;
        margin-bottom: 30px;
    }
    .stTextInput, .stTextArea {
        background-color: #f8f8f8;
        border-radius: 5px;
        border: 1px solid #008080;
        padding: 10px;
    }
    .stButton > button {
        background-color: #6a0dad;
        color: white;
        font-weight: bold;
        border-radius: 25px;
        padding: 10px 25px;
        font-size: 16px;
        border: none;
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        background-color: #4b0082;
        box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    }
    .step-header {
        background-color: #008080;
        color: white;
        padding: 10px;
        border-radius: 5px;
        margin-bottom: 20px;
    }
    .subheader {
        color: #6a0dad;
        font-size: 20px;
        font-weight: bold;
        margin-bottom: 10px;
    }
    [class="big-font"]{
        margin-top : -110px;
    }
    </style>
    """, unsafe_allow_html=True)

hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)
 
st.markdown('<p class="big-font">LEGAL DRAFT PRO</p>', unsafe_allow_html=True)
st.markdown('<p class="medium-font">Streamline Your Legal Document Creation</p>', unsafe_allow_html=True)

doc1_path = r"files/my_own.docx"
doc2_path = r"files/my_own2.docx"
doc3_path = r"files/Data_license_Agreement.docx"
doc4_path = r"files/professional_service_agreement.docx"
doc5_path = r"files/asset_purchase_agreement.docx"
doc6_path = r"files/SAFE2.docx"
doc7_path = r"files/Stock_Purchase_Agreement_Startups.docx"
 
doc1 = Document(doc1_path)
doc2 = Document(doc2_path)
doc3=Document(doc3_path)
doc4=Document(doc4_path)
doc5=Document(doc5_path)
doc6=Document(doc6_path)
doc7=Document(doc7_path)
 
placeholders1 = re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc1.paragraphs]))
placeholders2 = re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc2.paragraphs]))
placeholders3 = re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc3.paragraphs]))
placeholders4 = re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc4.paragraphs]))
placeholders5= re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc5.paragraphs]))
placeholders6= re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc6.paragraphs]))
placeholders7= re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc7.paragraphs]))
#Document path and loading to collect the Placeholders 
# doc_path = r'my_own.docx'  
# doc = Document(doc_path)
# placeholders = re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc.paragraphs]))

#Functions 
DOCUMENT_FIXED_LINES = {
    "Master Service Agreement": "2.SERVICES",
    "New York Agreement": "CONSULTANT UNDERSTANDS THAT THIS AGREEMENT AFFECTS IMPORTANT RIGHTS. BY SIGNING BELOW,",
    "Data License Agreement": "Parties agree as follows:",
    "Professional Service Agreement": "IN WITNESS WHEREOF, the parties hereto have executed this Agreement as of the date first above written.",
    "Asset Purchase Agreement": "PURCHASE AND SALE",
    "Safe Simple Agreement for Future Equity": "Definitions. Capitalized terms not otherwise defined in this SAFE will have the meanings ",
    "Founders Stock Purchase Agreement": "IN WITNESS WHEREOF, the parties have executed this Stock Purchase Agreement as of the date first above written."
}

# ... (previous code remains unchanged)
def process_input(user_input, placeholders1,placeholders2,placeholders3,placeholders4,placeholders5,placeholders6,placeholders7):
    template = """
You are an expert in filling details in word documents using placeholders and providing definitions for legal conditions.
Use the provided details to fill out the placeholders and provide definitions for the conditions mentioned.

Determine which document the details are for based on the context provided.
Indicate which document the details belong to by specifying one of the following at the beginning of your response:
- "Document: Master Service Agreement"
- "Document: New York Agreement"
- "Document: Data License Agreement"
- "Document: Professional Service Agreement"
- "Document: Asset Purchase Agreement"
- "Document: Safe Simple Agreement for Future Equity"
- "Document: Founders Stock Purchase Agreement"

Here are the placeholders for each document:
Document 1 (New York Agreement) placeholders: {placeholders1}
Document 2 (Master Service Agreement) placeholders: {placeholders2}
Document 3 (Data License Agreement) placeholders: {placeholders3}
Document 4 (Professional Service Agreement) placeholders: {placeholders4}
Document 5 (Asset Purchase Agreement) placeholders: {placeholders5}
Document 6 (Safe Simple Agreement for Future Equity) placeholders: {placeholders6}
Document 7 (Founders Stock Purchase Agreement) placeholders: {placeholders7}

For Master Service Agreement, please give only date, don't give the year and other matter.

For each placeholder, provide the information in this format:
PLACEHOLDER: information

If a placeholder is missing information, use this format:
PLACEHOLDER: MISSING

Based on the user input, fill in the placeholders and provide definitions for relevant legal terms.
Structure your response exactly like this, replacing the examples with actual content:
[[
    "document": "Master Service Agreement" or "New York Agreement" or "Data License Agreement" or "Professional Service Agreement" or "Asset Purchase Agreement" or "Safe Simple Agreement for Future Equity" or "Founders Stock Purchase Agreement",
    "placeholders": [
        [["PLACEHOLDER1": "Value1"]],
        [["PLACEHOLDER2": "Value2"]],
        [["PLACEHOLDER3": "MISSING"]]
    ],
    "definitions": [
        [["Term1": "Definition1"]],
        [["Term2": "Definition2"]]
    ]
]]

User input: {content}

Response:
"""
    formatted_template = template.format(placeholders1=placeholders1,
                                         placeholders2=placeholders2,
                                         placeholders3=placeholders3,
                                         placeholders4=placeholders4,
                                         placeholders5=placeholders5,
                                         placeholders6=placeholders6,
                                         placeholders7=placeholders7,
                                         content=user_input)
    prompt = PromptTemplate(template=formatted_template)
    chain = prompt | llm
    response = chain.invoke({"content": user_input})
    return response
# Process input function
def add_content_to_document(doc_path, placeholders, definitions, document_type):
    doc = Document(doc_path)
    fixed_line = DOCUMENT_FIXED_LINES.get(document_type, "Signatures:")
    
    # Replace placeholders
    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            if isinstance(value, str) and value != "MISSING":
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)
    
    # Add definitions before the fixed line
    for i, paragraph in enumerate(doc.paragraphs):
        if fixed_line in paragraph.text:
            for term, definition in definitions.items():
                doc.paragraphs[i].insert_paragraph_before("")
                
                # Insert the term and definition
                new_paragraph = doc.paragraphs[i].insert_paragraph_before(f"{term}: {definition}")
                
                # Apply blue color to the entire definition
                run = new_paragraph.runs[0]
                font = run.font
                font.color.rgb = RGBColor(0, 0, 255)  # Blue color (RGB)
                
                # Insert another empty paragraph
                doc.paragraphs[i].insert_paragraph_before("")
                # doc.paragraphs[i].insert_paragraph_before("")
                # doc.paragraphs[i].insert_paragraph_before(f"{term}: {definition}")
                # doc.paragraphs[i].insert_paragraph_before("")
            break
    return doc

# Initialize session state
if 'state' not in st.session_state:
    st.session_state.state = {
        'user_input': "",
        'collected_details': {},
        'definitions': [],
        'placeholders': [],
        'processed': False,
        'document_generated': False,
        'final_doc': None,
        'document_type': ""
    }

# User input text area
st.markdown('<div class="step-header">Step 1: Enter Your Query</div>', unsafe_allow_html=True)
user_input = st.text_area("Enter your query to fill the details:", 
                          value=st.session_state.state['user_input'])

# Process input
if user_input and st.button("Process Input"):
    global iii
    iii=0
    st.session_state.state['user_input'] = user_input
    with st.spinner("Processing your input..."):
        processed_response = process_input(user_input, placeholders1, placeholders2, placeholders3, placeholders4, placeholders5, placeholders6, placeholders7)
        processed_response = processed_response.content
        
        fresponse = processed_response.replace('[[', '{').replace(']]', '}')
        try:
            fresponse = fresponse.split("```json")[1].split("```")[0]
            fresponse = json.loads(fresponse)
        except:
            fresponse = json.loads(fresponse)
        
        st.session_state.state['document_type'] = fresponse.get("document", "")
        st.session_state.state['placeholders'] = fresponse.get("placeholders", [])
        st.session_state.state['definitions'] = fresponse.get("definitions", [])
        st.session_state.state['processed'] = True
    st.success("Input processed successfully!")

# Display placeholders and definitions
if st.session_state.state['processed']:
    st.markdown('<div class="step-header">Step 2: Review and Update Details</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)

    with col1:
        st.markdown('<p class="subheader">Missing Details</p>', unsafe_allow_html=True)
        for placeholder in st.session_state.state['placeholders']:
            for key, value in placeholder.items():
                if value == "MISSING":
                    user_detail = st.text_input(f"{key.replace('_', ' ')}:", key=iii,
                                                value=st.session_state.state['collected_details'].get(key, ""))
                    st.session_state.state['collected_details'][key] = user_detail
                else:
                    st.text_input(f"{key}:",key=iii, value=value, disabled=True)
                iii+1

    with col2:
        st.markdown('<p class="subheader">Definitions</p>', unsafe_allow_html=True)
        updated_definitions = []
        for definition in st.session_state.state['definitions']:
            updated_definition = {}
            for term, desc in definition.items():
                updated_desc = st.text_area(f"{term}:", value=desc, height=100, key=f"def_{term}")
                updated_definition[term] = updated_desc
            updated_definitions.append(updated_definition)
        
    if st.button("Update Details and Definitions"):
        updated_placeholders = []
        for placeholder in st.session_state.state['placeholders']:
            updated_placeholder = {}
            for key, value in placeholder.items():
                if value == "MISSING" and key in st.session_state.state['collected_details']:
                    updated_placeholder[key] = st.session_state.state['collected_details'][key]
                else:
                    updated_placeholder[key] = value
            updated_placeholders.append(updated_placeholder)
        
        st.session_state.state['placeholders'] = updated_placeholders
        st.session_state.state['definitions'] = updated_definitions
        st.success("Details and definitions updated successfully!")

    st.markdown('<div class="step-header">Step 3: Generate and Download Document</div>', unsafe_allow_html=True)
    if st.button("Generate Final Document"):
        with st.spinner("Generating document..."):
            doc_paths = {
                "New York Agreement": doc1_path,
                "Master Service Agreement": doc2_path,
                "Data License Agreement": doc3_path,
                "Professional Service Agreement": doc4_path,
                "Asset Purchase Agreement": doc5_path,
                "Safe Simple Agreement for Future Equity": doc6_path,
                "Founders Stock Purchase Agreement": doc7_path
            }
            document_type = st.session_state.state['document_type']
            doc_path = doc_paths.get(document_type)
            
            if doc_path:
                placeholders = {k: v for d in st.session_state.state['placeholders'] for k, v in d.items()}
                definitions = {k: v for d in st.session_state.state['definitions'] for k, v in d.items()}
                st.session_state.state['final_doc'] = add_content_to_document(
                    doc_path, 
                    placeholders, 
                    definitions,
                    document_type
                )
                st.session_state.state['document_generated'] = True
                st.success("Final document generated with all missing details and definitions.")
            else:
                st.error(f"No document path found for document type: {document_type}")

    if st.session_state.state['document_generated']:
        bio = io.BytesIO()
        st.session_state.state['final_doc'].save(bio)
        st.download_button(
            label="Download Final Document",
            data=bio.getvalue(),
            file_name="final_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Display the final content of placeholders and definitions
    if st.checkbox("Show final content"):
        # for x in st.session_state.state['placeholders']:
        #     st.json(x)
        st.json(st.session_state.state['placeholders'])
        st.json(st.session_state.state['definitions'])

else:
    st.info("Please enter your query and click 'Process Input' to start.")

st.markdown('</div>', unsafe_allow_html=True)